# -*- coding: utf-8 -*-
"""
Word 报告生成模块
生成差异代谢物药用筛选分析报告（.docx 格式）
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def make_branch_green(cell):
    """将表格单元格背景设为浅绿色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9F0D3')  # 浅绿色
    tcPr.append(shd)


def set_col_width(table, col_idx, width_inches):
    """设置表格指定列宽"""
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


def generate_word_report(
    annotated_df,
    pharma_df,
    star_df,
    abundance_df,
    params: dict,
    fig_heatmap_bytes=None,
    fig_boxplot_bytes=None,
    fig_vip_bytes=None,
    fig_bubble_bytes=None,
    fig_enr_bar_bytes=None,
    enr_df=None,
) -> bytes:
    """
    生成 Word 报告文档。

    Parameters
    ----------
    annotated_df : pd.DataFrame
        完整注释后的差异代谢物表
    pharma_df : pd.DataFrame
        候选药用化合物表
    star_df : pd.DataFrame
        明星分子评分表
    abundance_df : pd.DataFrame
        五物种丰度表
    params : dict
        分析参数字典 {vip_thresh, p_thresh, n_diff, n_pharma, n_star, ...}
    fig_heatmap_bytes : bytes, optional
        热图 PNG 字节流
    fig_boxplot_bytes : bytes, optional
        箱线图 PNG 字节流
    fig_vip_bytes : bytes, optional
        VIP 分布图 PNG 字节流
    fig_bubble_bytes : bytes, optional
        KEGG 气泡图 PNG 字节流
    fig_enr_bar_bytes : bytes, optional
        KEGG 柱状图 PNG 字节流
    enr_df : pd.DataFrame, optional
        KEGG 富集结果表

    Returns
    -------
    bytes
        Word 文档的字节流
    """
    doc = Document()

    # ==================== 页面设置 ====================
    section = doc.sections[0]
    section.page_width = Inches(11.69)   # A4 横向
    section.page_height = Inches(8.27)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ==================== 标题 ====================
    title = doc.add_heading('差异代谢物药用筛选分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 副标题信息
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').font.size = Pt(10)
    sub.add_run('　　').font.size = Pt(10)
    sub.add_run(f'分析工具：Metabolite Pharmaceutical Screener').font.size = Pt(10)

    doc.add_paragraph()

    # ==================== 一、分析概况 ====================
    h1 = doc.add_heading('一、分析概况', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 参数字段
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    headers = ['参数', '值', '参数', '值']
    for i, hdr in enumerate(headers):
        hdr_cells[i].text = hdr
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        make_branch_green(hdr_cells[i])

    vip_row = [f'VIP 阈值', f'> {params.get("vip_thresh", 1.0)}',
               f'P-value 阈值', f'< {params.get("p_thresh", 0.05)}']
    row2 = table.add_row().cells
    for i, val in enumerate(vip_row):
        row2[i].text = val

    n_diff = params.get('n_diff', len(annotated_df))
    n_named = params.get('n_named', 0)
    n_pharma = params.get('n_pharma', len(pharma_df))
    n_star = params.get('n_star', len(star_df))

    stat_row = ['差异代谢物总数', str(n_diff),
                '注释到化合物名', str(n_named)]
    row3 = table.add_row().cells
    for i, val in enumerate(stat_row):
        row3[i].text = val

    stat_row2 = ['候选药用化合物', str(n_pharma),
                  '明星分子', str(n_star)]
    row4 = table.add_row().cells
    for i, val in enumerate(stat_row2):
        row4[i].text = val

    doc.add_paragraph()

    # ==================== 二、差异代谢物筛选结果 ====================
    h2 = doc.add_heading('二、差异代谢物筛选结果', level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p = doc.add_paragraph(
        f'共筛选出 {n_diff} 个差异代谢物（VIP > {params.get("vip_thresh", 1.0)} 且 '
        f'P < {params.get("p_thresh", 0.05)}），其中 {n_named} 个成功匹配到化合物注释。'
    )

    # Top 20 差异代谢物表格
    doc.add_heading('Top 20 差异代谢物（按 VIP 排序）', level=2)
    top20 = annotated_df.nlargest(20, 'Vip_plsda')[
        ['Metabolite', 'Vip_plsda', 'P_value', 'FC', 'compound_name', 'super_class']
    ].reset_index(drop=True)
    top20.index = top20.index + 1
    top20.index.name = '排名'

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    col_names = ['代谢物名称', 'VIP (PLS-DA)', 'P-value', 'Fold Change', '化合物名', '化学分类']
    for i, c in enumerate(col_names):
        hdr[i].text = c
        hdr[i].paragraphs[0].runs[0].bold = True
        make_branch_green(hdr[i])

    for _, row_data in top20.iterrows():
        cells = tbl.add_row().cells
        cells[0].text = str(row_data['Metabolite'])[:40]
        cells[1].text = f"{row_data['Vip_plsda']:.4f}"
        cells[2].text = f"{row_data['P_value']:.2e}"
        cells[3].text = f"{row_data['FC']:.4f}"
        name = str(row_data.get('compound_name', '-'))
        cells[4].text = name[:30] if name != 'nan' else '-'
        sc = str(row_data.get('super_class', '-'))
        cells[5].text = sc[:25] if sc != 'nan' else '-'

    doc.add_paragraph()

    # ==================== 三、可视化 ====================
    h3 = doc.add_heading('三、可视化分析', level=1)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    if fig_heatmap_bytes:
        doc.add_heading('3.1 候选化合物丰度聚类热图（Top 30）', level=2)
        img_stream = io.BytesIO(fig_heatmap_bytes)
        doc.add_picture(img_stream, width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('图 1. 候选化合物丰度聚类热图（Z-score 标准化）')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic = True
        doc.add_paragraph()

    if fig_boxplot_bytes:
        doc.add_heading('3.2 物种丰度箱线图（Top 15）', level=2)
        img_stream = io.BytesIO(fig_boxplot_bytes)
        doc.add_picture(img_stream, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('图 2. 候选化合物在五种棘豆中的丰度分布')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic = True
        doc.add_paragraph()

    if fig_vip_bytes:
        doc.add_heading('3.3 VIP 分数分布', level=2)
        img_stream = io.BytesIO(fig_vip_bytes)
        doc.add_picture(img_stream, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('图 3. 差异代谢物 VIP 分数分布直方图')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic = True
        doc.add_paragraph()

    # ==================== 四、明星分子精选 ====================
    h4 = doc.add_heading('四、明星分子精选', level=1)
    for run in h4.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph(
        '综合评分公式：综合得分 = VIP得分 × 0.25 + 物种差异度 × 0.25 + '
        '药理类别 × 0.30 + 药理证据 × 0.20'
    )

    star_cols = ['Metabolite', 'compound_name', 'super_class', 'Vip_plsda',
                 '_diff_ratio', '_star_score', '_pharma_reason']
    available_cols = [c for c in star_cols if c in star_df.columns]
    top10_star = star_df.head(10)[available_cols].reset_index(drop=True)
    top10_star.index = top10_star.index + 1
    top10_star.index.name = '排名'

    tbl2 = doc.add_table(rows=1, cols=min(6, len(available_cols)))
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    display_names = {
        'Metabolite': '代谢物',
        'compound_name': '化合物名',
        'super_class': '化学分类',
        'Vip_plsda': 'VIP',
        '_diff_ratio': '差异倍数',
        '_star_score': '综合得分',
        '_pharma_reason': '筛选依据',
    }
    hdr2 = tbl2.rows[0].cells
    for i, col in enumerate(available_cols[:len(hdr2)]):
        hdr2[i].text = display_names.get(col, col)
        hdr2[i].paragraphs[0].runs[0].bold = True
        make_branch_green(hdr2[i])

    for _, row_data in top10_star.iterrows():
        cells = tbl2.add_row().cells
        for i, col in enumerate(available_cols[:len(cells)]):
            val = row_data.get(col, '-')
            if isinstance(val, float):
                cells[i].text = f'{val:.4f}'
            else:
                cells[i].text = str(val)[:30]

    doc.add_paragraph()

    # PubMed 链接列表
    doc.add_heading('Top 10 明星分子文献检索链接', level=2)
    for rank, (_, row) in enumerate(top10_star.iterrows(), 1):
        metab = row.get('Metabolite', '')
        compound = str(row.get('compound_name', '-'))
        search_name = compound if compound and compound != '-' else str(metab)
        if search_name in ('-', 'nan', ''):
            search_name = str(metab)
        pubmed_url = f'https://pubmed.ncbi.nlm.nih.gov/?term={search_name.replace(" ", "+")}'
        scholar_url = f'https://scholar.google.com/scholar?q={search_name.replace(" ", "+")}'
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{rank}. {search_name}：').bold = True
        p.add_run(f'PubMed: {pubmed_url}')
        p.add_run('　')
        p.add_run(f'Google Scholar: {scholar_url}')

    doc.add_paragraph()

    # ==================== 五、KEGG 通路富集分析 ====================
    if enr_df is not None and len(enr_df) > 0:
        h5 = doc.add_heading('五、KEGG 通路富集分析', level=1)
        for run in h5.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        doc.add_paragraph(
            f'使用超几何检验对 {len(pharma_df)} 个候选化合物进行 KEGG 通路富集分析，'
            f'共检测到 {len(enr_df[enr_df["P_value"] < 0.05])} 条显著富集通路（P < 0.05）。'
        )

        if fig_bubble_bytes:
            img_stream = io.BytesIO(fig_bubble_bytes)
            doc.add_picture(img_stream, width=Inches(6.0))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph('图 4. KEGG 通路富集气泡图')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].italic = True
            doc.add_paragraph()

        if fig_enr_bar_bytes:
            img_stream = io.BytesIO(fig_enr_bar_bytes)
            doc.add_picture(img_stream, width=Inches(5.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph('图 5. KEGG 通路富集柱状图（Top 20）')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].italic = True
            doc.add_paragraph()

        # 显著通路表格
        sig_enr = enr_df[enr_df['P_value'] < 0.05].sort_values('P_value').head(20)
        if len(sig_enr) > 0:
            doc.add_heading('显著富集通路详情（Top 20）', level=2)
            tbl3 = doc.add_table(rows=1, cols=5)
            tbl3.style = 'Table Grid'
            hdr3 = tbl3.rows[0].cells
            for i, h in enumerate(['KEGG ID', '通路名称', '命中数', 'Rich Factor', 'P-value']):
                hdr3[i].text = h
                hdr3[i].paragraphs[0].runs[0].bold = True
                make_branch_green(hdr3[i])

            for _, row_data in sig_enr.iterrows():
                cells = tbl3.add_row().cells
                cells[0].text = str(row_data.get('KEGG_ID', '-'))
                cells[1].text = str(row_data.get('Pathway_Name', '-'))[:40]
                cells[2].text = str(row_data.get('Candidate_Hits', '-'))
                cells[3].text = f"{row_data.get('Rich_Factor', 0):.3f}"
                cells[4].text = f"{row_data.get('P_value', 1):.2e}"

        doc.add_paragraph()

    # ==================== 六、分析方法说明 ====================
    h6 = doc.add_heading('六、分析方法', level=1)
    for run in h6.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    methods = [
        ('差异筛选', '采用 VIP（Variable Importance in Projection）> {vip_thresh} 且 P-value < {p_thresh} 进行差异代谢物筛选。'),
        ('药理分类', '基于化合物的 super_class（化学超类）字段，筛选 Alkaloids、Phenylpropanoids、Lignans、Flavonoids、Terpenoids 等潜在药用化学类别。'),
        ('丰度计算', '从原始 diff 文件中提取各物种的生物重复丰度值，计算算术平均值作为该物种的代表性丰度。'),
        ('明星分子评分', '综合 VIP 得分（0.3）、物种差异度（0.3）、药理类别（0.4）三个维度计算加权综合得分。'),
        ('KEGG 富集', '以所有注释到 KEGG 通路的差异代谢物为背景集，使用超几何检验（Fisher Exact Test）评估候选化合物的通路富集显著性，BH 法校正 P 值。'),
    ]
    for title_text, desc_text in methods:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title_text}：').bold = True
        p.add_run(desc_text.format(vip_thresh=params.get('vip_thresh', 1.0),
                                   p_thresh=params.get('p_thresh', 0.05)))

    # ==================== 尾部 ====================
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— 本报告由 Metabolite Pharmaceutical Screener 自动生成 —')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.italic = True

    # ==================== 输出 ====================
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
